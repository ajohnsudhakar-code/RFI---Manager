import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

import json
import os
import re
from datetime import date

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image as PILImage
except ImportError:
    print("ERROR: Pillow not installed. Run: pip install pillow")
    sys.exit(1)

# ── Stable module-level constant ─────────────────────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))

# ════════════════════════════════════════════════════════════
#  LAYOUT CONSTANTS  (A4 Landscape)
# ════════════════════════════════════════════════════════════
PAGE_W   = Inches(11.69)
PAGE_H   = Inches(8.27)
MARGIN_L = Inches(0.60)
MARGIN_R = Inches(0.60)
MARGIN_T = Inches(0.50)
MARGIN_B = Inches(0.50)
CW       = 10.49

C_ACCENT  = "1F497D"
C_INFOBG  = "D6E4F0"
C_CLIENTBG= "EBF3FB"
C_SIGBG   = "F0F4F8"
C_WHITE   = (0xFF, 0xFF, 0xFF)
C_RED     = RGBColor(0xC0, 0x00, 0x00)
C_BLUE    = RGBColor(0x1F, 0x49, 0x7D)
C_LINK    = RGBColor(0x00, 0x70, 0xC0)
C_GRAY    = RGBColor(0x66, 0x66, 0x66)
C_DARK    = RGBColor(0x44, 0x44, 0x44)

# ════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    ex = tcPr.find(qn('w:tcW'))
    if ex is not None:
        tcPr.remove(ex)
    tcPr.append(tcW)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_table_width(table, inches):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    ex = tblPr.find(qn('w:tblW'))
    if ex is not None:
        tblPr.remove(ex)
    tblPr.append(tblW)

def para(cell, text='', bold=False, size=10, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, new=False):
    p = cell.add_paragraph() if new else cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p

def spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)

def _fit_image(iw, ih, max_w, max_h):
    """Return (display_w, display_h) in inches maintaining aspect ratio."""
    if ih == 0:
        return max_w, max_h
    ratio = iw / ih
    w = min(max_w, max_h * ratio)
    h = w / ratio
    return w, h

def _get_img_info(snapshots):
    """Return list of (path, pixel_w, pixel_h) for each snapshot."""
    info = []
    for path in snapshots:
        try:
            with PILImage.open(path) as img:
                info.append((path, img.width, img.height))
        except Exception:
            info.append((path, 800, 600))
    return info

def _estimate_snap_height(snap_info, avail_w):
    """Estimate total vertical inches the snapshots will occupy."""
    n = len(snap_info)
    if n == 0:
        return 0.0
    cell_w = (avail_w - 0.2) / 2
    caption_h = 0.25
    if n == 1:
        _, iw, ih = snap_info[0]
        _, dh = _fit_image(iw, ih, avail_w, 6.0)
        return dh + caption_h
    if n == 2:
        _, iw0, ih0 = snap_info[0]
        _, iw1, ih1 = snap_info[1]
        if (iw0 >= ih0) and (iw1 >= ih1):
            _, dh0 = _fit_image(iw0, ih0, avail_w, 3.5)
            _, dh1 = _fit_image(iw1, ih1, avail_w, 3.5)
            return dh0 + dh1 + caption_h * 2
        else:
            _, dh0 = _fit_image(iw0, ih0, cell_w, 5.0)
            _, dh1 = _fit_image(iw1, ih1, cell_w, 5.0)
            return max(dh0, dh1) + caption_h
    total = 0.0
    for i in range(0, n, 2):
        row_h = 0.0
        for j in range(min(2, n - i)):
            _, iw, ih = snap_info[i + j]
            _, dh = _fit_image(iw, ih, cell_w, 3.5)
            row_h = max(row_h, dh)
        total += row_h + caption_h
    return total

def _borderless_table(doc, ncols, avail_w):
    """Add a borderless ncols-column table to doc."""
    tbl = doc.add_table(rows=1, cols=ncols)
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        tblB.append(el)
    tblPr.append(tblB)
    set_table_width(tbl, avail_w)
    return tbl

def _fill_img_cell(cell, path, dw):
    """Put image into a table cell."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_picture(path, width=Inches(dw))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def _embed_snapshots(doc, snap_info, avail_w):
    """Embed snapshots using smart layout (full-width / side-by-side / grid)."""
    n = len(snap_info)
    if n == 0:
        return
    cell_w = (avail_w - 0.2) / 2

    if n == 1:
        path, iw, ih = snap_info[0]
        dw, _ = _fit_image(iw, ih, avail_w, 5.0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(path, width=Inches(dw))
        return

    if n == 2:
        path0, iw0, ih0 = snap_info[0]
        path1, iw1, ih1 = snap_info[1]
        both_landscape = (iw0 >= ih0) and (iw1 >= ih1)
        if both_landscape:
            for path, iw, ih in [(path0, iw0, ih0), (path1, iw1, ih1)]:
                dw, _ = _fit_image(iw, ih, avail_w, 3.2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(0)
                p.add_run().add_picture(path, width=Inches(dw))
        else:
            tbl = _borderless_table(doc, 2, avail_w)
            c0, c1 = tbl.cell(0, 0), tbl.cell(0, 1)
            set_cell_width(c0, cell_w)
            set_cell_width(c1, cell_w)
            dw0, _ = _fit_image(iw0, ih0, cell_w - 0.1, 4.5)
            dw1, _ = _fit_image(iw1, ih1, cell_w - 0.1, 4.5)
            _fill_img_cell(c0, path0, dw0)
            _fill_img_cell(c1, path1, dw1)
        return

    for row_start in range(0, n, 2):
        row_items = snap_info[row_start:row_start + 2]
        ncols     = len(row_items)
        each_w    = avail_w / ncols
        tbl       = _borderless_table(doc, ncols, avail_w)
        for j, (path, iw, ih) in enumerate(row_items):
            cell = tbl.cell(0, j)
            set_cell_width(cell, each_w - 0.05)
            dw, _ = _fit_image(iw, ih, each_w - 0.15, 3.0)
            _fill_img_cell(cell, path, dw)


# ════════════════════════════════════════════════════════════
#  MAIN CALLABLE
# ════════════════════════════════════════════════════════════
def generate_rfi_document(config: dict) -> dict:
    """Generate the RFI Word document from a config dict.

    Returns a dict with keys:
        success      : bool
        output_path  : str  (empty on failure)
        message      : str
        logs         : list[str]
    """
    logs = []
    try:
        _BASE_EARLY = os.path.normpath(os.path.join(_HERE, ".."))

        # ── Extract config sections ───────────────────────────────────────────
        _project  = config.get("project", {})
        _co       = config.get("company", {})
        if not _co:
            _co = config.get("originator", {})
        _orig_cfg = config.get("originator", {})
        _client   = config.get("client", {})
        _paths    = config.get("paths", {})
        _settings = config.get("settings", {})
        _approved = config.get("approved_rfis", [])

        # ── Project / company / originator ───────────────────────────────────
        PROJECT_NAME    = _project.get("name", "")
        PROJECT_ADDRESS = _project.get("address", "")
        CONTRACT_NO     = _project.get("project_number", "")

        ORIGINATOR_NAME  = _orig_cfg.get("name",  _co.get("originator_name",  ""))
        ORIGINATOR_TITLE = _orig_cfg.get("title", _co.get("originator_title", ""))
        ORIGINATOR_PHONE = _orig_cfg.get("phone", _co.get("originator_phone", ""))
        ORIGINATOR_EMAIL = _orig_cfg.get("email", _co.get("originator_email", ""))
        COMPANY_NAME     = _co.get("name", _co.get("company_name", _co.get("company", "")))
        COMPANY_WEB      = _co.get("website", _co.get("company_website", ""))

        CLIENT_COMPANY = _client.get("company", "")
        CLIENT_ATTN    = _client.get("attn",    "")
        CLIENT_EMAIL   = _client.get("email",   "")
        CLIENT_PHONE   = _client.get("phone",   "")
        CLIENT_ROLE    = _client.get("role",    "")

        SNAPSHOTS_DIR = _paths.get("snapshots", "")
        OUTPUT_DIR    = _paths.get("output",    "")
        MAX_SNAPSHOTS = int(_settings.get("max_snapshots", 6))

        # ── Local helpers ─────────────────────────────────────────────────────
        def _abs(p):
            """Return absolute path; search SNAPSHOTS_DIR, _HERE, _BASE_EARLY."""
            if not p:
                return ""
            if os.path.isabs(p) and os.path.exists(p):
                return p
            if SNAPSHOTS_DIR:
                c = os.path.join(SNAPSHOTS_DIR, p)
                if os.path.exists(c):
                    return c
            return p

        def get_snapshots(rfi_num):
            snaps = []
            for i in range(1, MAX_SNAPSHOTS + 1):
                path = os.path.join(SNAPSHOTS_DIR, f"RFI_{rfi_num:03d}_snap{i}.png")
                if os.path.exists(path):
                    snaps.append(path)
            return snaps

        # ── Logo / signature ──────────────────────────────────────────────────
        LOGO_IMG      = _abs(_paths.get("logo",      _co.get("company_logo",    "")))
        SIGNATURE_IMG = _abs(_paths.get("signature", _co.get("signature_image", "")))

        if not LOGO_IMG or not os.path.exists(LOGO_IMG):
            LOGO_IMG = ""
        if not SIGNATURE_IMG or not os.path.exists(SIGNATURE_IMG):
            SIGNATURE_IMG = ""

        # ── Dates / filename ──────────────────────────────────────────────────
        today     = date.today()
        today_str = today.strftime("%d %B %Y")
        if len(_approved) == 1:
            _rfi0 = _approved[0].get('rfi_number', 1)
            if isinstance(_rfi0, str):
                _m = re.search(r'\d+', _rfi0)
                _rfi0 = int(_m.group()) if _m else 1
            filename = f"RFI_{_rfi0:03d}_{PROJECT_NAME.replace(' ', '_')}_{today.strftime('%d%m%Y')}.docx"
        else:
            filename = f"RFI_{PROJECT_NAME.replace(' ', '_')}_{today.strftime('%d%m%Y')}.docx"

        # ── Directories ───────────────────────────────────────────────────────
        if SNAPSHOTS_DIR:
            os.makedirs(SNAPSHOTS_DIR, exist_ok=True)
        if OUTPUT_DIR:
            os.makedirs(OUTPUT_DIR, exist_ok=True)

        # ── Output path ───────────────────────────────────────────────────────
        out_path = os.path.join(OUTPUT_DIR, filename)
        logs.append(f"Building {len(_approved)} RFIs - A4 Landscape")

        # ════════════════════════════════════════════════════════════
        #  DOCUMENT GENERATION
        # ════════════════════════════════════════════════════════════
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)

        for idx, issue in enumerate(_approved):
            rfi_num = issue.get('rfi_number', idx + 1)
            if isinstance(rfi_num, str):
                m = re.search(r'\d+', rfi_num)
                rfi_num = int(m.group()) if m else idx + 1
            rfi_ref    = f"RFI-{rfi_num:03d}"
            sheets_str  = issue.get('sheets', issue.get('pages', ''))
            priority    = issue.get('priority', '')
            response_by = issue.get('response_required_by', '')

            logs.append(f"  Building {rfi_ref}...")

            # ── Load snapshots early so we can pick orientation ───────────
            snapshots = get_snapshots(rfi_num)
            snap_info = _get_img_info(snapshots)

            est_h        = _estimate_snap_height(snap_info, 9.5)
            use_portrait = est_h > 4.5

            if use_portrait:
                pg_w   = Inches(8.27)
                pg_h   = Inches(11.69)
                cw     = 7.07
                orient = WD_ORIENT.PORTRAIT
                logs.append(f"    Portrait layout (estimated snap height {est_h:.1f}\")")
            else:
                pg_w   = PAGE_W
                pg_h   = PAGE_H
                cw     = CW
                orient = WD_ORIENT.LANDSCAPE

            s            = cw / CW
            snap_avail_w = round(cw - 0.4, 3)

            # ── Section setup ─────────────────────────────────────────────
            if idx == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section()

            section.orientation   = orient
            section.page_width    = pg_w
            section.page_height   = pg_h
            section.left_margin   = MARGIN_L
            section.right_margin  = MARGIN_R
            section.top_margin    = MARGIN_T
            section.bottom_margin = MARGIN_B

            # ── 1. HEADER TABLE ───────────────────────────────────────────
            hdr = doc.add_table(rows=1, cols=2)
            hdr.style = 'Table Grid'
            set_table_width(hdr, cw)

            lc_w = round(5.0 * s, 3)
            rc_w = round(cw - lc_w, 3)
            lc = hdr.cell(0, 0); rc = hdr.cell(0, 1)
            set_cell_width(lc, lc_w); set_cell_width(rc, rc_w)
            set_cell_margins(lc, 80, 80, 120, 120)
            set_cell_margins(rc, 80, 80, 120, 120)
            lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            lp = lc.paragraphs[0]
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after  = Pt(0)
            if LOGO_IMG and os.path.exists(LOGO_IMG):
                lp.add_run().add_picture(LOGO_IMG, height=Inches(0.5))
            lp2 = lc.add_paragraph()
            lp2.paragraph_format.space_before = Pt(2)
            lp2.paragraph_format.space_after  = Pt(0)
            r2 = lp2.add_run(COMPANY_NAME)
            r2.bold = True; r2.font.size = Pt(13)
            r2.font.color.rgb = C_RED
            if COMPANY_WEB:
                lp3 = lc.add_paragraph()
                lp3.paragraph_format.space_before = Pt(0)
                lp3.paragraph_format.space_after  = Pt(0)
                r3 = lp3.add_run(COMPANY_WEB)
                r3.font.size = Pt(9); r3.font.color.rgb = C_LINK

            first = True
            for label, value in [("RFI No:", rfi_ref), ("Date:", today_str),
                                  ("Project:", PROJECT_NAME), ("Contract No:", CONTRACT_NO)]:
                p = rc.paragraphs[0] if first else rc.add_paragraph()
                first = False
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                rl = p.add_run(f"{label}  "); rl.bold = True; rl.font.size = Pt(9)
                rv = p.add_run(value);        rv.font.size = Pt(9)

            spacer(doc, 4)

            # ── 2. TITLE ──────────────────────────────────────────────────
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tp.paragraph_format.space_before = Pt(2)
            tp.paragraph_format.space_after  = Pt(2)
            tr = tp.add_run(f"Request for Information  —  {rfi_ref}")
            tr.bold = True; tr.font.size = Pt(14)
            tr.font.color.rgb = C_BLUE

            spacer(doc, 4)

            # ── 3. ADDRESSEE TABLE ────────────────────────────────────────
            addr = doc.add_table(rows=2, cols=4)
            addr.style = 'Table Grid'
            set_table_width(addr, cw)

            col_a = round(1.5 * s, 3)
            col_b = round(4.0 * s, 3)
            col_c = round(1.5 * s, 3)
            col_d = round(cw - col_a - col_b - col_c, 3)
            col_w = [col_a, col_b, col_c, col_d]

            for c_idx, (label, value) in enumerate([("To:", CLIENT_COMPANY), ("Role:", CLIENT_ROLE)]):
                lc = addr.cell(0, c_idx * 2)
                vc = addr.cell(0, c_idx * 2 + 1)
                set_cell_width(lc, col_w[c_idx * 2])
                set_cell_width(vc, col_w[c_idx * 2 + 1])
                set_cell_margins(lc, 50, 50, 100, 100)
                set_cell_margins(vc, 50, 50, 100, 100)
                set_cell_bg(lc, C_INFOBG)
                lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                para(lc, label, bold=True, size=9)
                para(vc, value, size=9)

            attn_lc = addr.cell(1, 0)
            attn_vc = addr.cell(1, 1).merge(addr.cell(1, 3))
            for _p in attn_vc.paragraphs[1:]:
                _p._element.getparent().remove(_p._element)
            set_cell_width(attn_lc, col_a)
            set_cell_margins(attn_lc, 50, 50, 100, 100)
            set_cell_margins(attn_vc, 50, 50, 100, 100)
            set_cell_bg(attn_lc, C_INFOBG)
            attn_lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            attn_vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para(attn_lc, "Attn:", bold=True, size=9)
            para(attn_vc, CLIENT_ATTN, size=9)

            email_row = addr.add_row()
            ec_l = email_row.cells[0]
            ec_v = email_row.cells[1].merge(email_row.cells[3])
            for _p in ec_v.paragraphs[1:]:
                _p._element.getparent().remove(_p._element)
            set_cell_bg(ec_l, C_INFOBG)
            set_cell_margins(ec_l, 50, 50, 100, 100)
            set_cell_margins(ec_v, 50, 50, 100, 100)
            ec_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            ec_v.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para(ec_l, "Email:", bold=True, size=9)
            para(ec_v, CLIENT_EMAIL, size=9)

            spacer(doc, 4)

            # ── 4. INFO TABLE ─────────────────────────────────────────────
            info = doc.add_table(rows=5, cols=2)
            info.style = 'Table Grid'
            set_table_width(info, cw)

            info_lw = round(1.5 * s, 3)
            info_vw = round(cw - info_lw, 3)
            info_data = [
                ("Project / Address",    f"{PROJECT_NAME}  |  {PROJECT_ADDRESS}"),
                ("Subject",              issue['description']),
                ("Drawing Reference",    sheets_str),
                ("Priority",             priority),
                ("Response Required By", response_by),
            ]
            for r_idx, (label, value) in enumerate(info_data):
                lc = info.cell(r_idx, 0); vc = info.cell(r_idx, 1)
                set_cell_width(lc, info_lw); set_cell_width(vc, info_vw)
                set_cell_margins(lc, 60, 60, 100, 100)
                set_cell_margins(vc, 60, 60, 100, 100)
                lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(lc, C_INFOBG)
                para(lc, label, bold=True, size=9)
                para(vc, value, size=9)

            spacer(doc, 4)

            # ── 5. MAIN RFI TABLE ─────────────────────────────────────────
            main = doc.add_table(rows=2, cols=2)
            main.style = 'Table Grid'
            set_table_width(main, cw)

            item_w = round(0.6 * s, 3)
            desc_w = round(cw - item_w, 3)

            h0 = main.cell(0, 0); h1 = main.cell(0, 1)
            set_cell_width(h0, item_w); set_cell_width(h1, desc_w)
            set_cell_margins(h0, 60, 60, 100, 100)
            set_cell_margins(h1, 60, 60, 100, 100)
            set_cell_bg(h0, C_ACCENT); set_cell_bg(h1, C_ACCENT)
            hp0 = para(h0, "Item", bold=True, size=10)
            hp0.runs[0].font.color.rgb = RGBColor(*C_WHITE)
            hp1 = para(h1, "Description / Drawing Reference", bold=True, size=10)
            hp1.runs[0].font.color.rgb = RGBColor(*C_WHITE)

            ic = main.cell(1, 0); dc = main.cell(1, 1)
            set_cell_width(ic, item_w); set_cell_width(dc, desc_w)
            set_cell_margins(ic, 80, 80, 100, 100)
            set_cell_margins(dc, 80, 80, 100, 100)
            ic.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            dc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            para(ic, f"{rfi_num}.1", size=10)
            para(dc, issue['description'], size=10)

            rp = dc.add_paragraph()
            rp.paragraph_format.space_before = Pt(4)
            rp.paragraph_format.space_after  = Pt(4)
            rr = rp.add_run(f"Reason: {issue['reason']}")
            rr.italic = True; rr.font.size = Pt(9)
            rr.font.color.rgb = C_DARK

            if not snap_info:
                wp = dc.add_paragraph()
                wp.add_run("[No snapshots — run crop_viewer.py first]").font.size = Pt(9)
                logs.append(f"    WARNING: No snapshots found for {rfi_ref}")

            spacer(doc, 4)

            # ── Snapshots (smart layout) ──────────────────────────────────
            if snap_info:
                _embed_snapshots(doc, snap_info, snap_avail_w)
                for i, (path, iw, ih) in enumerate(snap_info, 1):
                    logs.append(f"    Snap {i}: {os.path.basename(path)}  ({iw}x{ih}px)")
                spacer(doc, 4)

            # ── 6. DISCLAIMER ─────────────────────────────────────────────
            disc = doc.add_paragraph()
            disc.paragraph_format.space_before = Pt(2)
            disc.paragraph_format.space_after  = Pt(2)
            dr = disc.add_run(
                "Any instruction given in this advice note may require prior "
                "approval from an authorised third party."
            )
            dr.italic = True; dr.font.size = Pt(8)
            dr.font.color.rgb = C_GRAY

            spacer(doc, 4)

            # ── 7. SIGNATURE TABLE ────────────────────────────────────────
            sig = doc.add_table(rows=2, cols=2)
            sig.style = 'Table Grid'
            set_table_width(sig, cw)

            orig_cell = sig.cell(0, 0).merge(sig.cell(0, 1))
            for _p in orig_cell.paragraphs[1:]:
                _p._element.getparent().remove(_p._element)
            set_cell_width(orig_cell, cw)
            set_cell_margins(orig_cell, 60, 60, 120, 120)
            set_cell_bg(orig_cell, C_SIGBG)

            para(orig_cell, f"Originator: {ORIGINATOR_NAME}  |  {ORIGINATOR_TITLE}", size=9)
            op2 = orig_cell.add_paragraph()
            op2.paragraph_format.space_before = Pt(0)
            op2.paragraph_format.space_after  = Pt(0)
            op2.add_run(f"Ph: {ORIGINATOR_PHONE}  |  Email: {ORIGINATOR_EMAIL}").font.size = Pt(9)

            sc   = sig.cell(1, 0); dc3 = sig.cell(1, 1)
            sc_w = round(5.0 * s, 3); dc3_w = round(cw - sc_w, 3)
            set_cell_width(sc, sc_w); set_cell_width(dc3, dc3_w)
            set_cell_margins(sc,  80, 80, 120, 120)
            set_cell_margins(dc3, 80, 80, 120, 120)
            sc.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
            dc3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            sp4 = para(sc, "Signed:  ", size=9)
            if SIGNATURE_IMG and os.path.exists(SIGNATURE_IMG):
                try:
                    sp4.add_run().add_picture(SIGNATURE_IMG, height=Inches(0.30))
                    logs.append("    Signature embedded")
                except Exception as e:
                    logs.append(f"    Signature warning: {e}")

            para(dc3, f"Date:  {today_str}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
            logs.append(f"  {rfi_ref} done - OK")

        # ── Save ──────────────────────────────────────────────────────────────
        doc.save(out_path)

        return {
            "success":     True,
            "output_path": out_path,
            "message":     f"SUCCESS - File: {out_path} - RFIs: {len(_approved)}",
            "logs":        logs,
        }

    except Exception as e:
        return {
            "success":     False,
            "output_path": "",
            "message":     f"ERROR: {type(e).__name__}: {e}",
            "logs":        logs,
        }
